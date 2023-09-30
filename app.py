import gradio as gr
import json
import os
import docx
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches, Pt, Cm, Mm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert
import pandas as pd

with open('pokemon.json', 'r') as f:
    pokemons = json.load(f)

GEN_RANGE = {
    "1세대": [1, 151],
    "2세대": [152, 251],
    "3세대": [252, 386],
    "4세대": [387, 493],
    "5세대": [494, 649],
    "6세대": [650, 721],
    "7세대": [722, 809],
    "8세대": [810, 905],
    "9세대": [906, 1017]
}

generation = gr.Dropdown(
            [f"{k}세대" for k in range(1, 10)], value="1세대", label="포켓몬 세대", info="원하는 포켓몬 세대를 선택하세요."
        )

download = gr.File(label="Download a file")
text = gr.DataFrame()

def write_docx(gen):
    filename = f'포켓몬{gen}.docx'

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    #changing the page margins
    margin = 1.27
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    document.styles['Normal'].font.name = 'NanumSquareRound'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumSquareRound')

    data_dict = []
    start, end = GEN_RANGE[gen]
    for k in range(start, end+1):
        name = pokemons[k-1]['name']
        number = pokemons[k-1]['number']
        types = pokemons[k-1]['types']
        image_path = pokemons[k-1]['image_path']

        data_dict.append(
            dict(이름=name, No=number, 타입='+'.join(types))
        )

        df = pd.DataFrame(data_dict)
        # Document
        table = document.add_table(rows=4, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = f"{number}"
        hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(50)
        hdr_cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        hdr_cells = table.rows[1].cells
        p = hdr_cells[0].add_paragraph()
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(image_path, width=Cm(14.5), height=Cm(14.5))
        r.add_break(docx.enum.text.WD_BREAK.LINE)

        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = f"{name}"
        hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(70)
        hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 192, 192)
        hdr_cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = f"{'+'.join(types)}"
        hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(70)
        hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 192, 192)
        hdr_cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_page_break()
        yield df[['No', '이름', '타입']], filename.replace('docx', 'pdf')
    
    if filename not in os.listdir():
        document.save(filename)
        convert(filename)
    return df, filename.replace('docx', 'pdf')

demo = gr.Interface(write_docx, generation, [text, download], title="대치동 포켓몬 도감 생성기",
                    description="원하는 포켓몬 세대를 선택하고, 다운로드를 눌러주세요.")
demo.queue(concurrency_count=3)
demo.launch(share=True)